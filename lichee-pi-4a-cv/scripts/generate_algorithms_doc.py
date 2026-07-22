#!/usr/bin/env python3
"""Generate SCANX algorithm documentation (Word)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT = Path(r"C:\Users\16092007\New-word.docx")


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph("")


def build() -> Document:
    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("SCANX")
    tr.bold = True
    tr.font.size = Pt(26)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Алгоритмы обработки видео\nСвязь с исходным кодом")
    sr.font.size = Pt(14)

    doc.add_page_break()

    heading(doc, "1. Общая схема")
    para(
        doc,
        "Программа построена как конвейер (pipeline). Каждый кадр проходит цепочку модулей. "
        "Точка входа — main.py, ядро — VideoPipeline в cv_module/pipeline/processor.py.",
    )
    code_block(
        doc,
        "main.py\n"
        "  └─ load_config() → AppConfig\n"
        "  └─ create_app() → FastAPI + VideoPipeline + EventStore\n"
        "       └─ фоновый поток _loop():\n"
        "            источник → предобработка → детекция → события → отрисовка → JPEG для UI",
    )

    heading(doc, "2. Источники кадров")
    heading(doc, "2.1. Видеофайл", 2)
    para(doc, "Файл: cv_module/input/video_file.py, класс VideoFileSource", bold=True)
    bullets(
        doc,
        [
            "OpenCV VideoCapture открывает файл по пути video_path",
            "read() возвращает следующий кадр BGR",
            "При конце файла и loop_video=true — перемотка на кадр 0",
            "describe() возвращает строку file:имя_файла",
        ],
    )

    heading(doc, "2.2. Камера", 2)
    para(doc, "Файл: cv_module/input/camera_source.py, класс CameraSource", bold=True)
    bullets(
        doc,
        [
            "Windows: DirectShow (CAP_DSHOW), затем запасной бэкенд",
            "Linux: V4L2",
            "open(): VideoCapture(camera_id), разрешение 640×480, чтение тестового кадра",
            "read(): повтор до 3 попыток при сбое",
            "probe(): перебор индексов 0–3 для списка /cameras",
        ],
    )
    para(doc, "Выбор камеры в конфиге (cv_module/config.py):", bold=True)
    bullets(
        doc,
        [
            "input_source=builtin → camera_id по умолчанию 0",
            "input_source=external → camera_id по умолчанию 1",
            "UI передаёт camera_id через POST /config",
        ],
    )

    heading(doc, "3. Предобработка кадров")
    para(doc, "Файл: cv_module/preprocessing/preprocessor.py, класс FramePreprocessor", bold=True)
    para(doc, "Алгоритм process(frame):", bold=True)
    bullets(
        doc,
        [
            "Счётчик кадров _index увеличивается на 1",
            "frame_skip: если >0, каждый (frame_skip+1)-й кадр пропускается → return None",
            "resize к input_size (по умолчанию 640×640), INTER_LINEAR",
            "BGR→RGB для blob (YOLO использует original BGR в detect)",
            "Возвращает PreprocessResult(original, resized, scale_x, scale_y)",
        ],
    )
    para(
        doc,
        "Важно: детектор получает prep.original — полный кадр исходного разрешения. "
        "YOLO внутри сам масштабирует через blobFromImage до input_size.",
    )

    heading(doc, "4. Гибридный детектор")
    para(doc, "Файл: cv_module/inference/detector.py, класс HybridDetector", bold=True)
    para(doc, "Метод detect(frame_bgr) объединяет два канала:", bold=True)
    code_block(
        doc,
        "yolo_dets = YoloOnnxDetector.detect(frame)     # person, vehicle\n"
        "heur_dets = HeuristicDetector.detect(frame)  # fire, smoke, water\n"
        "all = yolo_dets + heur_dets\n"
        "all = _dedupe(all)                           # слияние близких bbox\n"
        "return filter(confidence >= threshold)",
    )

    heading(doc, "4.1. Дедупликация HybridDetector._dedupe", 2)
    bullets(
        doc,
        [
            "Ключ: (object_class, bbox.x // 20, bbox.y // 20)",
            "В одной ячейке 20×20 px остаётся детекция с max confidence",
            "Предотвращает дубли от YOLO и эвристик на одном объекте",
        ],
    )

    heading(doc, "5. YOLOv8n ONNX — люди и транспорт")
    para(doc, "Файл: cv_module/inference/yolo_onnx.py, класс YoloOnnxDetector", bold=True)

    heading(doc, "5.1. Загрузка модели _load()", 2)
    bullets(
        doc,
        [
            "onnxruntime.InferenceSession(models/yolov8n.onnx)",
            "Провайдер: CPUExecutionProvider",
            "80 классов COCO; используются только person и VEHICLES",
        ],
    )

    heading(doc, "5.2. Инференс detect(frame_bgr)", 2)
    para(doc, "Шаги:", bold=True)
    bullets(
        doc,
        [
            "1. blobFromImage: нормализация /255, размер input_size, swapRB=True",
            "2. session.run → тензор предсказаний [1, 84, 8400] или аналог",
            "3. Транспонирование preds = out[0].T — строка на каждый anchor",
            "4. Для каждой строки: cx,cy,bw,bh = row[:4]; scores = row[4:]",
            "5. cid = argmax(scores); score = scores[cid]",
            "6. Отсечение score < confidence_threshold",
            "7. Пересчёт bbox из нормированных координат в пиксели кадра",
        ],
    )
    code_block(
        doc,
        "x1 = (cx - bw/2) * w / input_w\n"
        "y1 = (cy - bh/2) * h / input_h\n"
        "x2 = (cx + bw/2) * w / input_w\n"
        "y2 = (cy + bh/2) * h / input_h",
    )

    heading(doc, "5.3. NMS (Non-Maximum Suppression) _nms()", 2)
    bullets(
        doc,
        [
            "Сортировка bbox по score по убыванию",
            "Жадный алгоритм: берём лучший, удаляем пересекающиеся (IoU ≥ 0.45)",
            "IoU = intersection / union площадей прямоугольников",
        ],
    )

    heading(doc, "5.4. Маппинг классов _to_det()", 2)
    table(
        doc,
        ["COCO class", "→ object_class", "event_type", "severity"],
        [
            ["person", "person", "person_detected", "medium"],
            ["car, truck, bus, …", "vehicle", "vehicle_detected", "medium"],
            ["остальные 78 классов", "—", "отбрасываются", "—"],
        ],
    )

    heading(doc, "6. Эвристики — огонь, дым, вода")
    para(doc, "Файл: cv_module/inference/heuristics.py, класс HeuristicDetector", bold=True)
    para(
        doc,
        "YOLOv8n (COCO) не содержит классов fire/smoke/water. "
        "Эти объекты ищутся цветовым анализом в пространстве HSV.",
    )

    heading(doc, "6.1. Огонь и дым — _fire_smoke()", 2)
    para(doc, "Общий алгоритм для каждой маски:", bold=True)
    bullets(
        doc,
        [
            "1. cv2.cvtColor(BGR → HSV)",
            "2. cv2.inRange(hsv, lower, upper) → бинарная маска",
            "3. cv2.findContours(RETR_EXTERNAL) → контуры",
            "4. Фильтр: площадь контура ≥ 0.2% площади кадра",
            "5. cv2.boundingRect → bbox",
            "6. confidence = min(0.99, 0.5 + area_ratio)",
        ],
    )
    table(
        doc,
        ["Класс", "HSV lower", "HSV upper", "Смысл"],
        [
            ["fire", "(5, 120, 120)", "(35, 255, 255)", "Оранжево-красный, насыщенный, яркий"],
            ["smoke", "(0, 0, 80)", "(180, 60, 220)", "Низкая насыщенность, серые тона"],
        ],
    )

    heading(doc, "6.2. Вода — _water()", 2)
    bullets(
        doc,
        [
            "ROI: нижние 60% кадра (y ≥ 40% высоты) — вода обычно внизу",
            "HSV inRange (85,40,40)–(130,255,255) — синие/голубые оттенки",
            "MORPH_OPEN 5×5 — убрать шум",
            "Мин. площадь контура: 0.8% ROI",
            "confidence = min(0.95, 0.45 + area_ratio)",
            "bbox.y смещается на +40% высоты кадра (координаты полного кадра)",
        ],
    )

    heading(doc, "7. Конвейер обработки (live)")
    para(doc, "Файл: cv_module/pipeline/processor.py, класс VideoPipeline", bold=True)

    heading(doc, "7.1. Фоновый поток _loop()", 2)
    bullets(
        doc,
        [
            "Цикл while not stop:",
            "  ok, frame = source.read()",
            "  при ошибке — до 30 повторов с паузой 0.1с, иначе state=error",
            "  _process(frame, t0)",
            "При stop без pause — source.close()",
        ],
    )

    heading(doc, "7.2. Обработка кадра _process()", 2)
    bullets(
        doc,
        [
            "1. prep = preprocessor.process(frame); если None — кадр пропущен (frame_skip)",
            "2. t1 = time; dets = detector.detect(prep.original)",
            "3. latency_ms = скользящее среднее за 30 последних inference (мс)",
            "4. annotated = _draw(frame, dets) — рамки cv2.rectangle + подписи",
            "5. store.create() для каждой детекции",
            "6. fps = скользящее среднее 1/elapsed за 30 кадров",
            "7. _frame = annotated → отдаётся через GET /frame/latest как JPEG",
        ],
    )

    heading(doc, "7.3. Отрисовка _draw()", 2)
    para(doc, "Файл: cv_module/inference/classes.py — CLASS_COLORS, CLASS_LABELS", bold=True)
    bullets(
        doc,
        [
            "person — зелёный (0,200,0)",
            "vehicle — оранжевый (255,180,0)",
            "fire — красный (0,80,255)",
            "smoke — серый (180,180,180)",
            "water — жёлтый (255,200,0)",
        ],
    )

    heading(doc, "8. Система событий")
    para(doc, "Файл: cv_module/events/store.py, класс EventStore", bold=True)

    heading(doc, "8.1. Создание события create()", 2)
    bullets(
        doc,
        [
            "1. Проверка дедупликации: последние 20 событий",
            "   _same(bbox_a, bbox_b, tol=30): |Δx|,|Δy|,|Δw|,|Δh| < 30 px",
            "   Если тот же класс и похожий bbox → return None (не создавать)",
            "2. counter++; id = evt_{counter:06d}",
            "3. Опционально _save_frame: crop bbox+pad, resize max 320px, JPEG",
            "4. Event(...) с timestamp UTC, fps, detection_ms",
            "5. _trim(): если events > max_events (300) — удалить старейшие + jpg",
            "6. _save() → data/events.json",
        ],
    )

    heading(doc, "8.2. Формат events.json", 2)
    code_block(
        doc,
        '{\n  "counter": 42,\n  "events": [ { "id": "evt_000042", ... }, ... ]\n}',
    )

    heading(doc, "9. Алгоритм Render (полный проход)")
    para(doc, "Файл: cv_module/pipeline/render.py, функция render_video()", bold=True)
    bullets(
        doc,
        [
            "1. Открыть VideoCapture, прочитать fps, frame_count, width, height",
            "2. Создать HybridDetector и FramePreprocessor (как в live)",
            "3. Цикл cap.read() по всем кадрам (без пропуска конца файла)",
            "4. Для каждого кадра: prep = preprocessor.process(frame)",
            "   Если prep не None — detector.detect(prep.original)",
            "5. Каждая детекция → запись {frame, time_sec, time_code, bbox, ...}",
            "6. cap.release()",
            "7. Counter by_class → summary",
            "8. Вернуть JSON schema cv-module-render/v1 (не пишет на диск сервера)",
        ],
    )
    para(
        doc,
        "Отличие от live: render обрабатывает каждый кадр файла последовательно, "
        "не создаёт EventStore записи, не рисует на кадре, не ограничен max_events.",
    )

    heading(doc, "10. REST API и связь с алгоритмами")
    para(doc, "Файл: cv_module/api/app.py", bold=True)
    table(
        doc,
        ["Endpoint", "Вызывает в коде"],
        [
            ["POST /start", "pipeline.start() → _loop thread"],
            ["POST /pause", "pipeline.pause() → _stop.set(), state=PAUSED"],
            ["POST /stop", "pipeline.stop() → source.close()"],
            ["POST /config", "config.apply_updates(); pipeline.update_config(); _build_source()"],
            ["POST /render", "render_video(config, path, threshold)"],
            ["GET /frame/latest", "pipeline.get_frame_jpeg() — imencode JPEG q=90"],
            ["GET /events", "store.list_events(limit)"],
        ],
    )

    heading(doc, "11. Порог confidence_threshold")
    para(
        doc,
        "Единый порог из config/UI применяется в трёх местах:",
    )
    bullets(
        doc,
        [
            "YoloOnnxDetector: score < confidence → отброс до NMS",
            "HeuristicDetector: conf < confidence → отброс в detect()",
            "HybridDetector: финальный filter после _dedupe",
        ],
    )
    para(
        doc,
        "Ползунок UI (5–99%) → POST /config → detector.confidence обновляется "
        "в pipeline.update_config() для yolo и heuristics одновременно.",
    )

    heading(doc, "12. Mock-режим")
    para(doc, "Запуск: python main.py --mock", bold=True)
    bullets(
        doc,
        [
            "HybridDetector.detect() возвращает 3 фиксированных Detection",
            "Без загрузки ONNX, для тестирования UI и API",
        ],
    )

    heading(doc, "13. Карта файлов алгоритмов")
    table(
        doc,
        ["Модуль", "Алгоритм"],
        [
            ["inference/yolo_onnx.py", "YOLOv8n ONNX, NMS, COCO→person/vehicle"],
            ["inference/heuristics.py", "HSV маски fire/smoke/water"],
            ["inference/detector.py", "HybridDetector, merge, dedupe"],
            ["preprocessing/preprocessor.py", "frame_skip, resize"],
            ["pipeline/processor.py", "Live loop, metrics, draw"],
            ["pipeline/render.py", "Offline full-video scan"],
            ["events/store.py", "Event dedup, trim, JSON persist"],
            ["input/video_file.py", "File capture + loop"],
            ["input/camera_source.py", "Camera capture DSHOW/V4L2"],
        ],
    )

    heading(doc, "14. Диаграмма потока данных (live)")
    code_block(
        doc,
        "VideoFileSource / CameraSource\n"
        "        │ read()\n"
        "        ▼\n"
        "   BGR frame (H×W×3)\n"
        "        │\n"
        "        ▼\n"
        " FramePreprocessor.process()\n"
        "        │ (frame_skip?)\n"
        "        ▼\n"
        " HybridDetector.detect(original)\n"
        "   ├─ YoloOnnxDetector → person, vehicle\n"
        "   └─ HeuristicDetector → fire, smoke, water\n"
        "        │\n"
        "        ▼\n"
        " list[Detection]\n"
        "        ├─► EventStore.create() × N\n"
        "        └─► _draw() → annotated frame → /frame/latest",
    )

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("\n\nSCANX — описание алгоритмов\nC:\\Users\\16092007\\SCANX\\lichee-pi-4a-cv")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def main() -> int:
    doc = build()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

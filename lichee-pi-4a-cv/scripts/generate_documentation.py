#!/usr/bin/env python3
"""Generate SCANX application documentation (Word)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT = Path(r"C:\Users\16092007\Documentation.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph("")


def build() -> Document:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("SCANX")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Документация по приложению\nМодуль компьютерного зрения САТК")
    sr.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Версия API: 2.2.0  |  Платформа: Windows (ПК)")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    add_heading(doc, "1. Назначение программы")
    add_para(
        doc,
        "SCANX — программный модуль компьютерного зрения для анализа видеопотока "
        "в реальном времени. Программа обнаруживает объекты на видеофайлах и с камер, "
        "формирует журнал событий и предоставляет веб-интерфейс для управления и просмотра.",
    )
    add_para(doc, "Основной цикл работы:", bold=True)
    add_bullets(
        doc,
        [
            "Выбор источника: видеофайл, встроенная или внешняя камера",
            "Обработка кадров нейросетью YOLOv8n и цветовыми эвристиками",
            "Фиксация событий при обнаружении объектов",
            "Отображение результата в браузере и через REST API",
        ],
    )

    add_heading(doc, "2. Обнаруживаемые классы объектов", 2)
    add_table(
        doc,
        ["Класс", "Название в интерфейсе", "Метод", "Критичность"],
        [
            ["person", "Человек", "YOLOv8n ONNX", "Средняя"],
            ["vehicle", "Транспорт", "YOLOv8n ONNX", "Средняя"],
            ["fire", "Огонь", "HSV-эвристика", "Критическая"],
            ["smoke", "Дым", "HSV-эвристика", "Высокая"],
            ["water", "Вода", "HSV-эвристика", "Высокая"],
        ],
    )
    add_para(
        doc,
        "Модель YOLOv8n распознаёт людей и транспорт. Огонь, дым и вода определяются "
        "по цветовым характеристикам кадра (эвристики OpenCV).",
    )

    add_heading(doc, "3. Системные требования")
    add_bullets(
        doc,
        [
            "ОС: Windows 10/11",
            "Python 3.10 или новее (устанавливается автоматически через install.bat)",
            "Свободное место: ~500 МБ (окружение, модель, данные)",
            "Браузер: Chrome, Edge, Firefox (современный)",
            "Для камеры: рабочий драйвер USB-камеры или встроенной камеры",
        ],
    )

    add_heading(doc, "4. Установка и запуск")
    add_heading(doc, "4.1. Структура дистрибутива", 3)
    add_bullets(
        doc,
        [
            r"C:\Users\16092007\SCANX\Установщик\ — скрипты установки и запуска",
            r"C:\Users\16092007\SCANX\lichee-pi-4a-cv\ — файлы приложения",
            r"C:\Users\16092007\SCANX\lichee-pi-4a-cv\videos\ — папка для видеофайлов пользователя",
            r"C:\Users\16092007\SCANX\Инструкция.txt — краткая памятка",
        ],
    )

    add_heading(doc, "4.2. Первичная установка (один раз)", 3)
    add_bullets(
        doc,
        [
            "Запустите Установщик\\install.bat",
            "Дождитесь установки Python, зависимостей и модели YOLOv8n",
            "Положите видеофайлы (.mp4, .avi, .mkv и др.) в папку videos\\",
        ],
    )

    add_heading(doc, "4.3. Ежедневная работа", 3)
    add_table(
        doc,
        ["Действие", "Файл", "Описание"],
        [
            ["Запуск", "Установщик\\run.bat", "Старт сервера, открытие браузера http://127.0.0.1:8080"],
            ["Остановка", "Установщик\\stop.bat", "Принудительная остановка сервера"],
            ["", "Закрытие окна run.bat", "Сервер останавливается автоматически"],
            ["Очистка данных", "Установщик\\cleanup.bat", "Удаление журнала событий и логов (после остановки)"],
        ],
    )

    add_heading(doc, "5. Веб-интерфейс")
    add_para(doc, "Адрес: http://127.0.0.1:8080", bold=True)

    add_heading(doc, "5.1. Главный экран", 3)
    add_bullets(
        doc,
        [
            "Область видео — текущий кадр с рамками обнаруженных объектов",
            "FPS — скорость обработки",
            "Бейдж состояния: stop / running / paused / error",
            "Плейсхолдер до нажатия «Старт»",
            "Красное сообщение при ошибке (камера, файл и т.д.)",
        ],
    )

    add_heading(doc, "5.2. Панель управления", 3)
    add_table(
        doc,
        ["Элемент", "Назначение"],
        [
            ["Источник ▾", "Выбор: Видеофайл / Встроенная камера / Внешняя USB-камера"],
            ["Список видео + ↻", "Выбор файла из videos\\, обновление списка"],
            ["Список камер + ↻", "Выбор камеры 0/1/2; ↻ — поиск подключённых камер"],
            ["Точность (ползунок)", "Порог уверенности 5–99%. Выше — меньше ложных срабатываний"],
            ["Рендер", "Полный проход по видео и экспорт JSON (только для файла)"],
            ["Старт / Пауза", "Запуск, пауза и возобновление обработки"],
            ["События", "Открыть журнал событий"],
        ],
    )

    add_heading(doc, "5.3. Горячие клавиши", 3)
    add_table(
        doc,
        ["Клавиша", "Действие"],
        [
            ["Space", "Старт / Пауза / Возобновление"],
            ["Ctrl+Alt+R", "Рендер видео в JSON (режим файла)"],
        ],
    )

    add_heading(doc, "5.4. Режимы источника", 3)
    add_table(
        doc,
        ["Режим", "Описание", "Камера по умолчанию"],
        [
            ["Видеофайл", "Воспроизведение из videos\\ с зацикливанием", "—"],
            ["Встроенная камера", "Камера ноутбука", "Индекс 0"],
            ["Внешняя USB-камера", "USB-веб-камера", "Индекс 1"],
        ],
    )
    add_para(
        doc,
        "Если камера не работает, попробуйте другой номер в списке (Камера 0, 1 или 2) "
        "и нажмите ↻ для обновления списка.",
    )

    add_heading(doc, "5.5. Журнал событий", 3)
    add_bullets(
        doc,
        [
            "Открывается кнопкой «События» справа",
            "Показывает последние 50 событий",
            "Для каждого: класс, точность %, FPS, время, задержка детекции, ID события",
            "Обновляется автоматически каждые 2 секунды",
        ],
    )

    add_heading(doc, "6. Принцип работы обработки")
    add_para(
        doc,
        "После нажатия «Старт» запускается фоновый поток обработки кадров:",
    )
    add_bullets(
        doc,
        [
            "Чтение кадра из видеофайла или камеры OpenCV",
            "Предобработка: изменение размера для модели, пропуск кадров (frame_skip)",
            "Детекция: YOLOv8n (люди, транспорт) + эвристики (огонь, дым, вода)",
            "Отрисовка рамок и подписей на кадре",
            "Запись событий в журнал",
            "Обновление метрик FPS и задержки",
        ],
    )
    add_para(doc, "Состояния конвейера:", bold=True)
    add_bullets(
        doc,
        [
            "stopped — остановлен",
            "running — идёт обработка",
            "paused — пауза (видеофайл сохраняет позицию)",
            "error — ошибка (нет камеры, файл недоступен и т.д.)",
        ],
    )

    add_heading(doc, "7. Система событий")
    add_para(doc, "Каждое обнаружение создаёт событие со следующими полями:", bold=True)
    add_bullets(
        doc,
        [
            "id — evt_000001, evt_000002, …",
            "timestamp — время UTC",
            "object_class — person, vehicle, fire, smoke, water",
            "confidence — уверенность 0.0–1.0",
            "bbox — координаты рамки (x, y, width, height)",
            "fps, detection_ms — метрики в момент события",
        ],
    )
    add_para(doc, "Хранение:", bold=True)
    add_bullets(
        doc,
        [
            "Файл data\\events.json — журнал до 300 последних событий",
            "Сохранение кадров событий отключено по умолчанию (экономия места)",
            "Повторные срабатывания на том же месте подавляются (дедупликация)",
        ],
    )

    add_heading(doc, "8. Функция «Рендер»")
    add_para(
        doc,
        "Доступна только в режиме «Видеофайл». Выполняет полный проход по всему видео "
        "от начала до конца без пропуска кадров и формирует JSON-отчёт со всеми детекциями.",
    )
    add_para(doc, "Запуск:", bold=True)
    add_bullets(doc, ["Кнопка «Рендер»", "Горячая клавиша Ctrl+Alt+R"])
    add_para(doc, "Содержимое JSON-отчёта:", bold=True)
    add_bullets(
        doc,
        [
            "Информация о видео: FPS, число кадров, длительность, разрешение",
            "Список всех детекций: кадр, время, класс, точность, bbox",
            "Сводка: total_detections, by_class, by_label",
        ],
    )
    add_para(doc, "Файл сохраняется через диалог «Сохранить как» в браузере (имя: имя_видео_render.json).")

    add_heading(doc, "9. REST API")
    add_para(doc, "Базовый URL: http://127.0.0.1:8080", bold=True)

    add_heading(doc, "9.1. Основные методы", 3)
    add_table(
        doc,
        ["Метод", "URL", "Назначение"],
        [
            ["GET", "/health", "Проверка работоспособности"],
            ["GET", "/status", "Состояние, FPS, ошибки, последние детекции"],
            ["GET", "/videos", "Список видео в videos\\"],
            ["GET", "/cameras", "Список доступных камер"],
            ["GET", "/events?limit=N", "Журнал событий"],
            ["GET", "/events/{id}", "Одно событие по ID"],
            ["GET", "/frame/latest", "Последний кадр (JPEG)"],
            ["GET", "/metrics", "FPS, latency, счётчики"],
            ["POST", "/start", "Запуск / возобновление"],
            ["POST", "/stop", "Остановка"],
            ["POST", "/pause", "Пауза"],
            ["POST", "/config", "Обновление настроек"],
            ["POST", "/render", "Полный рендер видео в JSON"],
        ],
    )

    add_heading(doc, "9.2. POST /config — параметры", 3)
    add_table(
        doc,
        ["Параметр", "Тип", "Описание"],
        [
            ["input_source", "string", "file | builtin | external"],
            ["video_path", "string", "Полный путь к видеофайлу"],
            ["camera_id", "int", "Индекс камеры (0, 1, 2…)"],
            ["confidence_threshold", "float", "Порог 0.05–0.99"],
            ["frame_skip", "int", "Пропуск кадров между детекциями (0–10)"],
        ],
    )

    add_heading(doc, "9.3. Пример POST /config", 3)
    p = doc.add_paragraph()
    p.style = "No Spacing"
    code = p.add_run(
        '{\n'
        '  "input_source": "file",\n'
        '  "video_path": "C:\\\\...\\\\videos\\\\demo.mp4",\n'
        '  "confidence_threshold": 0.45\n'
        '}'
    )
    code.font.name = "Consolas"
    code.font.size = Pt(9)

    add_heading(doc, "10. Настройки приложения")
    add_para(doc, "Файл config\\default.yaml (основные параметры для ПК):", bold=True)
    add_table(
        doc,
        ["Параметр", "По умолчанию", "Описание"],
        [
            ["input_source", "file", "Источник: file, builtin, external"],
            ["videos_dir", "videos", "Папка с видео"],
            ["loop_video", "true", "Зацикливать видеофайл"],
            ["camera_id", "0", "Индекс камеры"],
            ["confidence_threshold", "0.45", "Порог уверенности"],
            ["api_port", "8080", "Порт веб-сервера"],
            ["max_events", "300", "Макс. событий в журнале"],
            ["save_event_frames", "false", "Сохранять кадры событий на диск"],
        ],
    )

    add_heading(doc, "11. Пользовательские папки и файлы")
    add_table(
        doc,
        ["Путь", "Назначение"],
        [
            ["videos\\", "Ваши видеофайлы для анализа"],
            ["models\\yolov8n.onnx", "Нейросеть (скачивается при установке)"],
            ["data\\events.json", "Журнал событий"],
            ["data\\logs\\cv_module.log", "Лог работы программы"],
            ["data\\scanx.pid", "PID процесса (пока сервер запущен)"],
            [".venv\\", "Python-окружение (не удалять)"],
        ],
    )

    add_heading(doc, "12. Решение типичных проблем")
    add_table(
        doc,
        ["Проблема", "Решение"],
        [
            ["install.bat не работает", "Установите Python 3.10+ с python.org (галочка Add to PATH), повторите install.bat"],
            ["Сервер не открывается", "Проверьте, что порт 8080 свободен; запустите stop.bat и снова run.bat"],
            ["Нет видео в списке", "Положите .mp4 в videos\\, нажмите ↻"],
            ["Камера не работает", "Выберите другой номер камеры; нажмите ↻; закройте другие программы с камерой"],
            ["Много ложных срабатываний", "Поднимите ползунок «Точность» до 55–70%"],
            ["Мало детекций", "Опустите «Точность» до 30–40%"],
            ["Диск заполнился", "Запустите cleanup.bat после stop.bat"],
            ["Старый интерфейс", "Обновите страницу Ctrl+F5"],
        ],
    )

    add_heading(doc, "13. Запуск из командной строки (для разработчиков)")
    add_bullets(
        doc,
        [
            r"cd C:\Users\16092007\SCANX\lichee-pi-4a-cv",
            r".\.venv\Scripts\python.exe main.py",
            r".\.venv\Scripts\python.exe main.py --mock  — тестовые детекции без модели",
            r".\.venv\Scripts\python.exe main.py --port 9000  — другой порт",
        ],
    )

    add_heading(doc, "14. Ограничения")
    add_bullets(
        doc,
        [
            "Огонь, дым и вода определяются эвристиками — возможны ложные срабатывания на похожие цвета",
            "YOLOv8n — лёгкая модель; точность ниже, чем у крупных моделей",
            "Рендер большого видео может занять продолжительное время",
            "Документ описывает версию для Windows; развёртывание на встроенных платах не рассматривается",
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("\n\nSCANX — модуль компьютерного зрения САТК\nДокумент сгенерирован автоматически")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def main() -> int:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
